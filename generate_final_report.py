from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_detailed_academic_report():
    doc = Document()
    
    # --- Helper Functions ---
    def add_capture_placeholder(context, filepath):
        p = doc.add_paragraph()
        run = p.add_run(f"[CAPTURE DE CODE À INSÉRER]")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p2 = doc.add_paragraph()
        p2.add_run(f"Type de code : {context}\n").bold = True
        p2.add_run(f"Fichier concerné : {filepath}")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph() # Spacer

    def setup_styles():
        styles = doc.styles
        # Title style customization if needed
        
    # --- Title Page ---
    title = doc.add_heading('RAPPORT DE PROJET ACADÉMIQUE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    subtitle = doc.add_heading('Plateforme de Business Intelligence Prédictive', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2 = doc.add_paragraph('Analyse des ventes E-commerce (Online Retail)')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 1. Introduction générale ---
    doc.add_heading('1. Introduction générale', level=1)
    
    doc.add_heading('1.1 Contexte de la Business Intelligence', level=2)
    doc.add_paragraph(
        "Dans l'ère actuelle du Big Data, la Business Intelligence (BI) est devenue un levier stratégique "
        "incontournable pour les entreprises. Elle permet de transformer des données brutes en informations "
        "actionnables, facilitant ainsi la prise de décision. Ce projet s'inscrit dans cette dynamique en "
        "proposant une solution moderne, capable de traiter des flux transactionnels sans la lourdeur des "
        "infrastructures traditionnelles."
    )

    doc.add_heading('1.2 Problématique et enjeux décisionnels', level=2)
    doc.add_paragraph(
        "Les entreprises d'e-commerce génèrent des volumes massifs de données. La problématique principale "
        "est de réduire la latence entre la collecte de la donnée (achat) et son analyse (dashboard). "
        "L'enjeu est double : fournir une vue descriptive instantanée (KPIs) et anticiper les tendances futures "
        "(Prédictif) pour optimiser les stocks et la stratégie marketing."
    )

    doc.add_heading('1.3 Objectifs du projet', level=2)
    doc.add_paragraph(
        "L'objectif principal est de développer une plateforme full-stack intégrant la chaîne de valeur complète de la donnée :\n"
        "1. Extraction et transformation performante (ETL in-memory).\n"
        "2. Expostion via une API REST rapide.\n"
        "3. Restitution visuelle et interactive.\n"
        "4. Enrichissement par des modèles d'Intelligence Artificielle (Segmentation et Prévision)."
    )

    doc.add_heading('1.4 Méthodologie globale adoptée', level=2)
    doc.add_paragraph(
        "Nous avons adopté une approche 'Code-First' et 'Zero-DB' pour ce prototype académique. "
        "L'absence de persistance lourde (type SQL) permet une itération rapide sur les algorithmes "
        "et une performance maximale grâce au traitement en RAM (Pandas)."
    )

    # --- 2. Présentation du dataset ---
    doc.add_heading('2. Présentation et compréhension du dataset', level=1)
    
    doc.add_heading('2.1 Source et description du dataset', level=2)
    doc.add_paragraph(
        "Le jeu de données 'Online Retail', provenant de l'UCI Machine Learning Repository (id=352), "
        "contient les transactions d'un grossiste e-commerce basé au Royaume-Uni entre 2010 et 2011. "
        "Il représente un cas d'école idéal avec environ 540 000 lignes."
    )

    doc.add_heading('2.2 Structure des données', level=2)
    doc.add_paragraph(
        "Les principaux attributs sont :\n"
        "- InvoiceNo : Identifiant unique de la transaction.\n"
        "- StockCode : Référence produit.\n"
        "- Description : Libellé du produit.\n"
        "- Quantity & UnitPrice : Métriques quantitatives.\n"
        "- CustomerID : Clé de segmentation client."
    )

    doc.add_heading('2.3 Problèmes de qualité des données', level=2)
    doc.add_paragraph(
        "L'exploration a révélé plusieurs problèmes nécessitant un nettoyage :\n"
        "- Valeurs manquantes sur les CustomerID (critique pour la segmentation).\n"
        "- Transactions d'annulation (InvoiceNo commençant par 'C').\n"
        "- Prix unitaires parfois à zéro ou négatifs (erreurs de saisie)."
    )

    doc.add_heading('2.4 Justification du choix du dataset', level=2)
    doc.add_paragraph(
        "Ce dataset a été choisi pour sa richesse multidimensionnelle (Temps, Produit, Client, Géographie), "
        "permettant de déployer l'ensemble des techniques BI et ML visées par ce projet."
    )

    add_capture_placeholder("Exploration initiale des données (chargement, colonnes)", "backend/app/data_provider.py")

    # --- 3. Architecture ---
    doc.add_heading('3. Architecture et conception du système', level=1)
    
    doc.add_heading('3.1 Vue globale de l’architecture', level=2)
    doc.add_paragraph(
        "Le système suit une architecture 3-tiers modulaire :\n"
        "1. Data Layer : Module Python autonome (DataProvider) gérant l'ETL et le stockage en mémoire.\n"
        "2. API Layer : Serveur FastAPI exposant les ressources via HTTP.\n"
        "3. Presentation Layer : Application Next.js (React) pour le rendu utilisateur."
    )

    doc.add_heading('3.2 Choix technologiques', level=2)
    doc.add_paragraph(
        "- Backend : FastAPI (Python) pour sa rapidité et sa typage strict (Pydantic).\n"
        "- Traitement : Pandas pour la manipulation vectorisée de données.\n"
        "- Frontend : Next.js pour le SSR (Server-Side Rendering) et la performance.\n"
        "- ML : Scikit-learn et XGBoost, standards de l'industrie."
    )

    doc.add_heading('3.3 Justification de l’absence de base de données', level=2)
    doc.add_paragraph(
        "Pour un volume de 500k lignes, le chargement intégral en RAM est plus performant que des I/O disques "
        "vers une base SQL. Cela simplifie le déploiement et élimine la latence de requête (Overhead ORM)."
    )

    doc.add_heading('3.4 Organisation du projet', level=2)
    doc.add_paragraph(
        "Le projet est structuré en deux dossiers racines 'backend' et 'frontend', favorisant une séparation "
        "claire des responsabilités (Separation of Concerns)."
    )

    add_capture_placeholder("Configuration générale du backend (FastAPI App)", "backend/app/main.py")

    # --- 4. ETL ---
    doc.add_heading('4. Pipeline ETL en mémoire', level=1)
    
    doc.add_heading('4.1 Extraction des données', level=2)
    doc.add_paragraph(
        "L'extraction se fait via la librairie `ucimlrepo` ou depuis un cache local Parquet pour accélérer "
        "les démarrages successifs."
    )

    doc.add_heading('4.2 Exploration automatique', level=2)
    doc.add_paragraph(
        "Un mécanisme d'auto-détection mappe dynamiquement les noms de colonnes du fichier source vers "
        "notre schéma interne normalisé, rendant le système résilient aux changements de format."
    )

    doc.add_heading('4.3 Nettoyage et préparation', level=2)
    doc.add_paragraph(
        "Les lignes sans CustomerID sont écartées. Les quantités négatives sont flaggées comme 'Retours' "
        "via une colonne booléenne `is_return`."
    )

    doc.add_heading('4.4 Transformations', level=2)
    doc.add_paragraph(
        "Création de champs calculés : `TotalPrice` (Qté * Prix), `MonthStr` (Série temporelle), "
        "et conversion des types datetime."
    )

    doc.add_heading('4.5 Validation du pipeline', level=2)
    doc.add_paragraph(
        "Le pipeline implémente un pattern Singleton Thread-Safe pour garantir qu'une seule instance "
        "de données est chargée en mémoire, accessible par tous les threads de l'API."
    )

    add_capture_placeholder("Fonctions ETL principales (load_data, transform)", "backend/app/data_provider.py")

    # --- 5. BI ---
    doc.add_heading('5. Modélisation BI et indicateurs', level=1)
    
    doc.add_heading('5.1 Définition des indicateurs clés (KPIs)', level=2)
    doc.add_paragraph(
        "Les KPIs implémentés sont :\n"
        "- Chiffre d'Affaires Total (Sum TotalPrice)\n"
        "- Nombre de Commandes Uniques\n"
        "- Panier Moyen (CA / Nb Commandes)\n"
        "- Taux de Retour (Nb Retours / Nb Total)"
    )

    doc.add_heading('5.2 et 5.3 Agrégations et Analyses', level=2)
    doc.add_paragraph(
        "Les agrégations sont réalisées via `groupby` de Pandas sur les axes temporels (Mois) et produits "
        "(Top 10 Ventes). Cette approche permet des temps de réponse de l'ordre de la milliseconde."
    )
    
    doc.add_heading('5.4 Apport décisionnel', level=2)
    doc.add_paragraph("Ces indicateurs offrent une vue macro-économique immédiate de la santé de l'activité.")

    add_capture_placeholder("Calcul des KPIs côté backend", "backend/app/analytics.py")

    # --- 6. ML Segmentation ---
    doc.add_heading('6. Segmentation clients par Machine Learning', level=1)
    
    doc.add_heading('6.1 Présentation de l’approche RFM', level=2)
    doc.add_paragraph(
        "La segmentation RFM classe les clients selon leur Récence d'achat, leur Fréquence et le Montant total dépensé. "
        "C'est une métrique standard en marketing relationnel."
    )

    doc.add_heading('6.3 Algorithme K-Means', level=2)
    doc.add_paragraph(
        "Nous utilisons l'algorithme non-supervisé K-Means (Scikit-Learn) sur les scores RFM normalisés "
        "pour identifier statistiquement des groupes de comportements similaires (Clusters)."
    )

    doc.add_heading('6.4 Analyse des segments', level=2)
    doc.add_paragraph(
        "Le modèle génère typiquement 4 segments, allant des 'Clients Fidèles à haute valeur' aux "
        "'Clients perdus ou inactifs'."
    )

    add_capture_placeholder("Implémentation du modèle de segmentation (RFM + K-Means)", "backend/app/ml_models.py")

    # --- 7. Forecasting ---
    doc.add_heading('7. Prévision du chiffre d’affaires', level=1)
    
    doc.add_heading('7.2 et 7.3 Modèles (Simple vs XGBoost)', level=2)
    doc.add_paragraph(
        "Nous avons implémenté deux approches comparatives :\n"
        "1. Régression Linéaire simple : Baseline rapide mais limitée.\n"
        "2. XGBoost Regressor : Modèle d'ensemble (Gradient Boosting) performant."
    )

    doc.add_heading('7.4 Évaluation', level=2)
    doc.add_paragraph(
        "Le modèle XGBoost, entraîné sur des features temporelles enrichies (Mois, Trimestre, Index), "
        "surpasse la baseline en réduisant l'erreur absolue moyenne (MAE) et en capturant la saisonnalité."
    )

    add_capture_placeholder("Code du modèle XGBoost (Entraînement/Prédiction)", "backend/app/ml_models.py")

    # --- 8. API ---
    doc.add_heading('8. Implémentation de l’API FastAPI', level=1)
    
    doc.add_paragraph(
        "L'API sert de passerelle unique. Elle est documentée automatiquement et valide les entrées/sorties "
        "via Pydantic."
    )
    
    doc.add_heading('8.2 et 8.3 Endpoints BI et ML', level=2)
    doc.add_paragraph(
        "- GET /api/v1/kpis/overview : Indicateurs globaux.\n"
        "- POST /api/v1/ml/predict/xgboost-forecast : Obtenir les prévisions futures."
    )

    add_capture_placeholder("Déclaration des routes FastAPI", "backend/app/api/endpoints.py")

    # --- 9. Frontend ---
    doc.add_heading('9. Interface utilisateur Next.js', level=1)
    
    doc.add_heading('9.1 Architecture Frontend', level=2)
    doc.add_paragraph(
        "L'application React utilise des hooks (useEffect) pour interroger l'API de manière asynchrone "
        "et gérer les états de chargement."
    )

    doc.add_heading('9.3 Visualisation', level=2)
    doc.add_paragraph(
        "La librairie Recharts est utilisée pour tracer les courbes de prévision et les diagrammes en camembert "
        "des segments, offrant une interactivité (Tooltips, Légendes) native."
    )

    add_capture_placeholder("Appels fetch vers l'API (Logique Frontend)", "frontend/src/app/ml/page.tsx")

    # --- 10. Résultats ---
    doc.add_heading('10. Résultats et discussion', level=1)
    doc.add_paragraph(
        "L'intégration complète du pipeline démontre une fluidité exceptionnelle. Le dashboard se charge en "
        "quelques secondes lors du premier démarrage (ETL), puis instantanément. "
        "L'apport du Machine Learning permet de passer d'une analyse descriptive (ce qui s'est passé) "
        "à une analyse prescriptive (ce qui va se passer)."
    )

    # --- 11. Conclusion ---
    doc.add_heading('11. Conclusion et perspectives', level=1)
    doc.add_paragraph(
        "Ce projet a permis de valider la faisabilité technique d'une BI moderne 'Zero-DB'. "
        "Pour aller plus loin, l'ajout d'une authentification utilisateur et la persistance des modèles ML "
        "seraient les prochaines étapes logiques."
    )

    # Save
    report_path = os.path.join(os.getcwd(), 'Rapport_Academique_Projet_BI_Final.docx')
    doc.save(report_path)
    print(f"Rapport généré : {report_path}")

if __name__ == "__main__":
    create_detailed_academic_report()
