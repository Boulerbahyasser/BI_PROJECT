from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Rapport de Projet : Plateforme BI Prédictive E-commerce', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Expert : Antigravity (Advanced Agentic AI)')
    doc.add_paragraph('Dataset : UCI Online Retail (540,000 transactions)')
    
    # 1. Cahier de charges
    doc.add_heading('1. Cahier de charges', level=1)
    doc.add_paragraph(
        "L'objectif est de concevoir et réaliser une plateforme décisionnelle (BI) capable de traiter un "
        "volume massif de données transactionnelles en temps réel, sans base de données persistante, "
        "en utilisant une architecture in-memory haute performance."
    )
    
    doc.add_heading('Besoins Fonctionnels :', level=2)
    capabilities = [
        "Dashboard en temps réel avec indicateurs clés de performance (KPIs).",
        "Pipeline ETL (Extract, Transform, Load) robuste et résilient.",
        "Segmentation client avancée (Modèle RFM + K-Means).",
        "Prévision des ventes par Machine Learning (Régression Linéaire & XGBoost)."
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
        
    doc.add_heading('Contraintes Techniques :', level=2)
    doc.add_paragraph("Architecture in-memory (Pandas Profiling).", style='List Bullet')
    doc.add_paragraph("Front-end : Next.js 15+ avec Recharts.", style='List Bullet')
    doc.add_paragraph("Back-end : FastAPI avec chargement asynchrone.", style='List Bullet')
    doc.add_paragraph("Zéro base de données (Calculs directs sur DataFrame).", style='List Bullet')

    # 2. Architecture & Étapes de Réalisation
    doc.add_heading('2. Architecture & Étapes de Réalisation', level=1)
    
    doc.add_heading('Étape A : Pipeline ETL Robuste', level=2)
    doc.add_paragraph(
        "Fichier : backend/app/data_provider.py\n"
        "Phase 1 (Explore) : Détection dynamique des colonnes via mappage de similarité.\n"
        "Phase 2 (Transform) : Nettoyage des IDs clients, calcul des prix totaux, filtrage des retours.\n"
        "Phase 3 (Load) : Mise en cache locale (Parquet) et Singleton in-memory."
    )

    doc.add_heading('Étape B : Analyse Descriptive (BI)', level=2)
    doc.add_paragraph(
        "Fichier : backend/app/analytics.py\n"
        "Calcul des KPIs (Chiffre d'affaires, Panier Moyen, Taux de Retour) avec filtrage spatio-temporel performant."
    )

    doc.add_heading('Étape C : Intelligence Artificielle (ML)', level=2)
    doc.add_paragraph(
        "Fichier : backend/app/ml_models.py\n"
        "Segmentation RFM : Clustering K-Means pour identifier les clients 'Champions' et 'À Risque'.\n"
        "Prévision XGBoost : Modèle de Gradient Boosting avec features temporelles (Année, Mois, Trimestre)."
    )

    # 3. Focus XGBoost
    doc.add_heading('3. Analyse Prédictive avec XGBoost', level=1)
    doc.add_paragraph(
        "Pourquoi XGBoost ?\n"
        "XGBoost a été choisi pour sa capacité supérieure à capturer les tendances non-linéaires "
        "et les saisonnalités par rapport à une régression simple. Sa robustesse face aux données manquantes "
        "et sa rapidité d'exécution le rendent idéal pour une application BI décisionnelle."
    )
    
    # 4. Interface & Navigation (Mapping)
    doc.add_heading('4. Interface & Navigation (Mapping)', level=1)
    
    mapping = [
        ("Dashboard Principal", "frontend/src/app/page.tsx", "KPIs globaux et Top Produits."),
        ("Machine Learning", "frontend/src/app/ml/page.tsx", "Comparaison Forecasting XGBoost & Segments RFM.")
    ]
    
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Vue / Écran'
    hdr_cells[1].text = 'Path du Fichier'
    hdr_cells[2].text = 'Description'
    
    for view, path, desc in mapping:
        row_cells = table.add_row().cells
        row_cells[0].text = view
        row_cells[1].text = path
        row_cells[2].text = desc

    # Save
    report_path = os.path.join(os.getcwd(), 'Rapport_BI_Predictif.docx')
    doc.save(report_path)
    print(f"Rapport généré avec succès : {report_path}")

if __name__ == "__main__":
    create_report()
